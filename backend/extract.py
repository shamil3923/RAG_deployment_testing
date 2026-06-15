"""Turn any uploaded file into plain text.

- PDF  -> PyMuPDF text; pages with little/no text are rendered and read by a vision model.
- DOCX -> paragraphs + tables.
- Image-> described/OCR'd by a vision model.
- else -> decoded as UTF-8 text.
"""
import base64
import io

import fitz  # PyMuPDF
import docx

import config
import llm

IMAGE_EXTS = (".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp")

VISION_PROMPT = (
    "Extract ALL text from this image verbatim. Then briefly describe any charts, "
    "diagrams, tables, or visual content. If there is no text, just describe what you see."
)


def _vision_read(image_png: bytes) -> str:
    """Send a PNG image to the vision model and return extracted text + description."""
    b64 = base64.b64encode(image_png).decode()
    resp = llm.client.chat.completions.create(
        model=config.VISION_MODEL,
        messages=[{
            "role": "user",
            "content": [
                {"type": "text", "text": VISION_PROMPT},
                {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{b64}"}},
            ],
        }],
        temperature=0.0,
        max_tokens=1024,
    )
    return resp.choices[0].message.content.strip()


def _prep_image(raw: bytes) -> bytes:
    """Normalise to PNG and downscale large images to keep the vision request small."""
    from PIL import Image

    img = Image.open(io.BytesIO(raw)).convert("RGB")
    img.thumbnail((1536, 1536))
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def _extract_pdf(raw: bytes) -> str:
    doc = fitz.open(stream=raw, filetype="pdf")
    parts = []
    for page in doc:
        text = page.get_text().strip()
        if len(text) < 20:  # likely a scanned / image-only page -> use vision
            pix = page.get_pixmap(dpi=150)
            text = _vision_read(pix.tobytes("png"))
        parts.append(text)
    return "\n\n".join(parts)


def _extract_docx(raw: bytes) -> str:
    d = docx.Document(io.BytesIO(raw))
    parts = [p.text for p in d.paragraphs if p.text.strip()]
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_text(raw: bytes, filename: str) -> str:
    """Dispatch on file extension and return plain text."""
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return _extract_pdf(raw)
    if name.endswith(".docx"):
        return _extract_docx(raw)
    if name.endswith(IMAGE_EXTS):
        return _vision_read(_prep_image(raw))
    return raw.decode("utf-8", errors="ignore")
